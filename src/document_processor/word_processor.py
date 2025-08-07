"""
Word document processor
Converts Word (.docx) files to ProcessedDocument objects with text chunks
"""

import logging
import time
from pathlib import Path
from typing import List, Union, Tuple

from docx import Document
from docx.oxml.exceptions import InvalidXmlError
from docx.opc.exceptions import PackageNotFoundError

from src.models.document import (
    DocumentType,
    ProcessingStatus,
    DocumentChunk,
    DocumentMetadata,
    ProcessedDocument,
    ProcessingConfig
)


logger = logging.getLogger(__name__)


class WordProcessingError(Exception):
    """Raised when Word processing fails"""
    pass


class WordProcessor:
    """
    Processes Word (.docx) files into structured document chunks

    Converts document paragraphs into text chunks that can be embedded
    and retrieved by RAG systems.
    """

    def __init__(self, config: ProcessingConfig = None):
        """
        Initialize Word processor with configuration

        Args:
            config: Processing configuration, defaults to ProcessingConfig()
        """
        self.config = config or ProcessingConfig()
        self.config.validate_config()

        logger.info(f"WordProcessor initialized with max_chunk_size={self.config.max_chunk_size}")

    def process(self, file_path: Union[str, Path]) -> ProcessedDocument:
        """
        Process a Word file into a ProcessedDocument

        Args:
            file_path: Path to the Word file

        Returns:
            ProcessedDocument with chunks and metadata

        Raises:
            WordProcessingError: If processing fails
        """
        start_time = time.time()
        file_path = Path(file_path)

        try:
            logger.info(f"Starting Word processing: {file_path}")

            # Validate file
            self._validate_file(file_path)

            # Read Word document
            doc, paragraphs = self._read_word_file(file_path)

            # Create metadata
            metadata = self._create_document_metadata(file_path, doc, paragraphs)
            metadata.processing_status = ProcessingStatus.PROCESSING

            # Process paragraphs into chunks
            chunks = self._create_chunks_from_paragraphs(paragraphs, file_path.name)

            logger.debug(f"Paragraph processing results: {len(paragraphs)} paragraphs -> {len(chunks)} chunks")

            # Update metadata with results
            processing_time = time.time() - start_time
            metadata.total_chunks = len(chunks)
            metadata.processing_time_seconds = processing_time
            metadata.processing_status = ProcessingStatus.COMPLETED

            logger.info(f"Word processing completed: {len(chunks)} chunks in {processing_time:.2f}s")

            return ProcessedDocument(metadata=metadata, chunks=chunks)

        except Exception as e:
            processing_time = time.time() - start_time
            error_msg = f"Word processing failed for {file_path}: {str(e)}"
            logger.error(error_msg)

            # Create failed metadata if possible
            try:
                metadata = DocumentMetadata(
                    filename=file_path.name,
                    document_type=DocumentType.WORD,
                    file_size=file_path.stat().st_size if file_path.exists() else 0,
                    processing_status=ProcessingStatus.FAILED,
                    processing_time_seconds=processing_time,
                    error_message=str(e)
                )
                return ProcessedDocument(metadata=metadata, chunks=[])
            except:
                # If even metadata creation fails, raise the original error
                raise WordProcessingError(error_msg) from e

    def _validate_file(self, file_path: Path) -> None:
        """
        Validate that the file exists and is processable

        Args:
            file_path: Path to validate

        Raises:
            WordProcessingError: If validation fails
        """
        if not file_path.exists():
            raise WordProcessingError(f"File does not exist: {file_path}")

        if not file_path.is_file():
            raise WordProcessingError(f"Path is not a file: {file_path}")

        if file_path.stat().st_size == 0:
            raise WordProcessingError(f"File is empty: {file_path}")

        # Check file extension
        allowed_extensions = {'.docx'}
        if file_path.suffix.lower() not in allowed_extensions:
            raise WordProcessingError(
                f"Unsupported file extension '{file_path.suffix}'. "
                f"Supported: {allowed_extensions}"
            )

        # Check if file is readable
        try:
            with open(file_path, 'rb') as f:
                f.read(1024)  # Try to read first 1KB
        except PermissionError:
            raise WordProcessingError(f"Permission denied reading file: {file_path}")
        except Exception as e:
            raise WordProcessingError(f"Cannot read file {file_path}: {e}")

    def _read_word_file(self, file_path: Path) -> Tuple[Document, List[Tuple[str, str, int]]]:
        """
        Read Word file and extract paragraphs with metadata

        Args:
            file_path: Path to Word file

        Returns:
            Tuple of (Document object, List of (paragraph_text, paragraph_type, paragraph_index))

        Raises:
            WordProcessingError: If reading fails
        """
        try:
            # Read Word document
            doc = Document(file_path)

            # Extract paragraphs with metadata
            paragraphs = []
            current_heading = ""

            for i, paragraph in enumerate(doc.paragraphs):
                # Skip completely empty paragraphs if configured to do so
                if self.config.skip_empty_paragraphs and not paragraph.text.strip():
                    continue

                # Determine paragraph type
                para_type = self._get_paragraph_type(paragraph)

                # Track current heading for context
                if para_type.startswith("heading"):
                    current_heading = paragraph.text.strip()

                # Create paragraph info
                para_info = (
                    paragraph.text.strip(),
                    para_type,
                    i,
                    current_heading  # Add heading context
                )
                paragraphs.append(para_info)

            if not paragraphs:
                raise WordProcessingError(f"No readable paragraphs found in document: {file_path}")

            logger.debug(f"Successfully extracted {len(paragraphs)} paragraphs")
            return doc, paragraphs

        except (PackageNotFoundError, InvalidXmlError) as e:
            raise WordProcessingError(f"File appears to be corrupted or not a valid Word document: {e}")
        except Exception as e:
            if isinstance(e, WordProcessingError):
                raise
            raise WordProcessingError(f"Failed to read Word file: {e}")

    def _get_paragraph_type(self, paragraph) -> str:
        """
        Determine the type of paragraph (heading, normal, list, etc.)

        Args:
            paragraph: python-docx Paragraph object

        Returns:
            String describing paragraph type
        """
        style_name = paragraph.style.name.lower()

        if 'heading' in style_name:
            if 'heading 1' in style_name:
                return "heading_1"
            elif 'heading 2' in style_name:
                return "heading_2"
            elif 'heading 3' in style_name:
                return "heading_3"
            else:
                return "heading_other"
        elif 'list' in style_name or 'bullet' in style_name:
            return "list_item"
        elif 'quote' in style_name:
            return "quote"
        elif 'title' in style_name:
            return "title"
        else:
            return "normal"

    def _split_long_paragraph(self, text: str, max_length: int) -> List[str]:
        """
        Split very long paragraphs into contextual chunks with overlap

        DESIGN DECISION: Instead of creating single-sentence chunks, we create
        multi-sentence chunks with overlap to preserve context for better RAG retrieval.
        Each chunk should contain 2-3 sentences when possible for optimal context-to-precision ratio.

        Args:
            text: Text to potentially split
            max_length: Maximum allowed length for embedding compatibility

        Returns:
            List of overlapping text segments with preserved context
        """
        if len(text) <= max_length:
            return [text]

        # Split at sentence boundaries (periods followed by space and capital letter)
        import re
        sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)

        if len(sentences) <= self.config.target_sentences_per_chunk:
            # If we have few sentences, try word-level splitting as last resort
            return self._split_at_word_boundaries(text, max_length)

        segments = []
        i = 0

        while i < len(sentences):
            # Build chunk starting from sentence i
            current_chunk = sentences[i]
            sentences_in_chunk = 1

            # Add more sentences to reach target size or sentence count
            j = i + 1
            while (j < len(sentences) and
                   sentences_in_chunk < self.config.target_sentences_per_chunk and
                   len(current_chunk + " " + sentences[j]) <= max_length):
                current_chunk += " " + sentences[j]
                sentences_in_chunk += 1
                j += 1

            segments.append(current_chunk.strip())

            # Move forward, but with overlap for context preservation
            overlap_sentences = min(self.config.overlap_sentences, sentences_in_chunk - 1)
            i = j - overlap_sentences  # Step back for overlap

            # Prevent infinite loops
            if i <= len(segments) - 1:
                i += 1

        return segments

    def _split_at_word_boundaries(self, text: str, max_length: int) -> List[str]:
        """
        Split text at word boundaries when sentence splitting isn't sufficient

        Args:
            text: Text to split
            max_length: Maximum segment length

        Returns:
            List of word-boundary split segments with overlap
        """
        words = text.split()
        if len(words) <= 3:  # Too few words to split meaningfully
            return [text]

        segments = []
        i = 0

        while i < len(words):
            current_segment = words[i]
            j = i + 1

            # Add words until we approach the limit
            while j < len(words) and len(current_segment + " " + words[j]) <= max_length:
                current_segment += " " + words[j]
                j += 1

            segments.append(current_segment)

            # Move forward with word-level overlap
            overlap_words = min(3, j - i - 1)  # Overlap 2-3 words for context
            i = j - overlap_words

            # Prevent infinite loops
            if i <= len(segments) - 1:
                i += 1

        return segments
        """
        Convert a paragraph to enhanced natural language text
        
        DESIGN DECISION: We use a naive natural language approach as a first step.
        Alternative approaches to consider with more data and performance analysis:
        - Structured document representation preserving exact formatting
        - Template-based conversion with document structure awareness
        - Semantic paragraph classification and importance weighting  
        - Cross-reference and citation preservation for technical documents
        
        Example:
            Input:
                para_text: "TechCorp supports flexible working arrangements and remote work options."
                para_type: "normal"
                current_heading: "Work Schedule and Remote Work Policy"
            
            Output:
                "Under Work Schedule and Remote Work Policy: TechCorp supports flexible working arrangements and remote work options."
        
        Args:
            para_text: Raw paragraph text
            para_type: Type of paragraph (heading, normal, list_item, etc.)
            current_heading: Current section heading for context
            
        Returns:
            Enhanced natural language description
        """
        if not para_text.strip():
            return ""

        # Clean the text
        clean_text = " ".join(para_text.split())  # Normalize whitespace

        # Add context based on paragraph type
        if para_type.startswith("heading"):
            # Headings are important structure - include level info
            level = para_type.replace("heading_", "").replace("_", " ")
            return f"Section {level}: {clean_text}"

        elif para_type == "title":
            return f"Document title: {clean_text}"

        elif para_type == "list_item":
            if current_heading:
                return f"Under {current_heading}, item: {clean_text}"
            else:
                return f"List item: {clean_text}"

        elif para_type == "quote":
            if current_heading:
                return f"Quote in {current_heading}: {clean_text}"
            else:
                return f"Quote: {clean_text}"

        else:  # normal paragraphs
            if current_heading:
                return f"Under {current_heading}: {clean_text}"
            else:
                return clean_text

    def _create_chunks_from_paragraphs(
        self,
        paragraphs: List[Tuple[str, str, int, str]],
        source: str
    ) -> List[DocumentChunk]:
        """
        Create document chunks from paragraphs with contextual overlap

        DESIGN DECISION: Create chunks of 2-3 sentences (200-800 chars) with overlap
        to optimize for RAG retrieval. Larger chunks provide better context for understanding,
        while overlap ensures semantic continuity across chunk boundaries.

        Args:
            paragraphs: List of (text, type, index, current_heading) tuples
            source: Source filename

        Returns:
            List of DocumentChunk objects with optimal size and context
        """
        chunks = []

        # Convert paragraphs to enhanced text
        paragraph_texts = []
        for para_text, para_type, para_idx, current_heading in paragraphs:
            try:
                enhanced_text = self._convert_paragraph_to_text(para_text, para_type, current_heading)
                if enhanced_text and enhanced_text.strip():  # More lenient check
                    paragraph_texts.append((enhanced_text, para_type, para_idx, current_heading))
                else:
                    logger.debug(f"Empty enhanced text for paragraph {para_idx}: '{para_text[:50]}...'")
            except Exception as e:
                logger.error(f"Failed to convert paragraph {para_idx}: {e}")
                # Include original text as fallback
                fallback_text = para_text.strip() if para_text else f"Paragraph {para_idx}"
                if fallback_text:
                    paragraph_texts.append((fallback_text, para_type, para_idx, current_heading))

        if not paragraph_texts:
            error_msg = f"No valid paragraphs extracted from {source}. Input had {len(paragraphs)} paragraphs."
            logger.error(error_msg)
            raise WordProcessingError(error_msg)  # Break instead of silent failure

        # Process paragraphs into contextual chunks
        all_segments = []  # List of (text, metadata) for all segments

        for enhanced_text, para_type, para_idx, heading in paragraph_texts:
            # Handle very long paragraphs by splitting them if needed
            if self.config.enforce_embedding_limits and len(enhanced_text) > self.config.max_embedding_chars:
                # Split the paragraph into embedding-compatible segments with overlap
                segments = self._split_long_paragraph(enhanced_text, self.config.max_embedding_chars)
                logger.info(f"Split long paragraph {para_idx} into {len(segments)} segments")

                for segment_idx, segment in enumerate(segments):
                    segment_meta = {
                        "text": segment,
                        "para_type": para_type,
                        "para_idx": f"{para_idx}_seg{segment_idx}",
                        "heading": heading
                    }
                    all_segments.append(segment_meta)
            else:
                # Normal paragraph - add as single segment
                segment_meta = {
                    "text": enhanced_text,
                    "para_type": para_type,
                    "para_idx": para_idx,
                    "heading": heading
                }
                all_segments.append(segment_meta)

        logger.debug(f"Created {len(all_segments)} segments from {len(paragraph_texts)} paragraphs")

        # If we have no segments, return early
        if not all_segments:
            logger.warning(f"No segments created from {source}")
            return chunks

        # Group segments into optimal-sized chunks with overlap
        chunk_index = 0
        i = 0

        while i < len(all_segments):
            # Build current chunk
            current_chunk_text = ""
            current_chunk_paragraphs = []
            current_chunk_types = set()
            current_chunk_headings = set()
            segments_in_chunk = 0

            # Add segments until we reach good chunk size
            j = i
            while j < len(all_segments):
                segment = all_segments[j]
                potential_text = current_chunk_text + "\n" + segment["text"] if current_chunk_text else segment["text"]

                # Check if we should break (but always include at least one segment)
                should_break = False
                if segments_in_chunk > 0:  # Only break if we have at least one segment
                    # Break if we'd exceed max chunk size
                    if len(potential_text) > self.config.max_chunk_size:
                        should_break = True

                    # Prefer to start new chunks with headings (but only if current chunk is substantial)
                    elif (segment["para_type"].startswith("heading") and
                          len(current_chunk_text) >= self.config.min_chunk_size):
                        should_break = True

                if should_break:
                    break

                # Add this segment to current chunk
                current_chunk_text = potential_text
                current_chunk_paragraphs.append(segment["para_idx"])
                current_chunk_types.add(segment["para_type"])
                if segment["heading"]:
                    current_chunk_headings.add(segment["heading"])

                segments_in_chunk += 1
                j += 1

            # Create chunk if we have content
            if current_chunk_text and current_chunk_text.strip():
                chunk = self._create_chunk(
                    content=current_chunk_text,
                    source=source,
                    chunk_index=chunk_index,
                    paragraph_indices=current_chunk_paragraphs,
                    paragraph_types=list(current_chunk_types),
                    headings=list(current_chunk_headings)
                )
                chunks.append(chunk)
                chunk_index += 1
                logger.debug(f"Created chunk {chunk_index} with {segments_in_chunk} segments, {len(current_chunk_text)} chars")

            # Move to next position with simple overlap
            if j >= len(all_segments):
                break

            # Simple overlap: go back 1 segment for context (if we used more than 1)
            overlap_amount = min(1, segments_in_chunk - 1) if segments_in_chunk > 1 else 0
            i = j - overlap_amount

            # Ensure we always make progress
            if i <= chunk_index - 1:  # Prevent infinite loops
                i = j

        logger.info(f"Created {len(chunks)} chunks from {len(all_segments)} segments")

        # Fallback: if no chunks were created but we have segments, create simple chunks
        if not chunks and all_segments:
            logger.warning("No chunks created with complex logic, falling back to simple chunking")

            current_text = ""
            current_paragraphs = []
            fallback_chunk_index = 0

            for segment in all_segments:
                potential_text = current_text + "\n" + segment["text"] if current_text else segment["text"]

                if len(potential_text) <= self.config.max_chunk_size:
                    current_text = potential_text
                    current_paragraphs.append(segment["para_idx"])
                else:
                    # Save current chunk if it has content
                    if current_text.strip():
                        fallback_chunk = self._create_chunk(
                            content=current_text,
                            source=source,
                            chunk_index=fallback_chunk_index,
                            paragraph_indices=current_paragraphs,
                            paragraph_types=["normal"],
                            headings=[]
                        )
                        chunks.append(fallback_chunk)
                        fallback_chunk_index += 1

                    # Start new chunk
                    current_text = segment["text"]
                    current_paragraphs = [segment["para_idx"]]

            # Add final chunk
            if current_text.strip():
                fallback_chunk = self._create_chunk(
                    content=current_text,
                    source=source,
                    chunk_index=fallback_chunk_index,
                    paragraph_indices=current_paragraphs,
                    paragraph_types=["normal"],
                    headings=[]
                )
                chunks.append(fallback_chunk)

            logger.info(f"Fallback created {len(chunks)} simple chunks")

        return chunks

        return chunks

    def _create_chunk(
        self,
        content: str,
        source: str,
        chunk_index: int,
        paragraph_indices: List[int],
        paragraph_types: List[str],
        headings: List[str]
    ) -> DocumentChunk:
        """
        Create a single DocumentChunk with proper metadata

        Args:
            content: Text content of the chunk
            source: Source filename
            chunk_index: Index of this chunk in the document
            paragraph_indices: List of original paragraph indices
            paragraph_types: List of paragraph types in this chunk
            headings: List of section headings relevant to this chunk

        Returns:
            DocumentChunk object
        """
        chunk_id = f"{Path(source).stem}_para_{chunk_index:03d}"

        # Count sentences for context evaluation
        import re
        sentence_count = len(re.findall(r'[.!?]+', content))

        metadata = {
            "paragraph_indices": paragraph_indices,
            "paragraph_count": len(paragraph_indices),
            "paragraph_types": paragraph_types,
            "section_headings": headings,
            "character_count": len(content),
            "sentence_count": sentence_count,
            "has_headings": any(ptype.startswith("heading") for ptype in paragraph_types),
            "has_overlap": chunk_index > 0,  # All chunks after first may have overlap
            "chunk_quality": "good" if self.config.min_chunk_size <= len(content) <= self.config.max_chunk_size else "size_warning"
        }

        return DocumentChunk(
            chunk_id=chunk_id,
            content=content.strip(),
            source=source,
            chunk_index=chunk_index,
            metadata=metadata
        )

    def _convert_paragraph_to_text(self, para_text: str, para_type: str, current_heading: str) -> str:
        """Convert a paragraph to enhanced natural language text"""

        if not para_text or not para_text.strip():
            return ""

        # Clean the text
        clean_text = " ".join(para_text.split())  # Normalize whitespace

        # Add context based on paragraph type
        if para_type.startswith("heading"):
            level = para_type.replace("heading_", "").replace("_", " ")
            enhanced_text = f"Section {level}: {clean_text}"
        elif para_type == "title":
            enhanced_text = f"Document title: {clean_text}"
        elif para_type == "list_item":
            if current_heading:
                enhanced_text = f"Under {current_heading}, item: {clean_text}"
            else:
                enhanced_text = f"List item: {clean_text}"
        elif para_type == "quote":
            if current_heading:
                enhanced_text = f"Quote in {current_heading}: {clean_text}"
            else:
                enhanced_text = f"Quote: {clean_text}"
        else:  # normal paragraphs
            if current_heading:
                enhanced_text = f"Under {current_heading}: {clean_text}"
            else:
                enhanced_text = clean_text

        return enhanced_text
    def _create_document_metadata(
        self,
        file_path: Path,
        doc: Document,
        paragraphs: List[Tuple[str, str, int, str]]
    ) -> DocumentMetadata:
        """
        Create document metadata from file and document information

        Args:
            file_path: Path to the Word file
            doc: python-docx Document object
            paragraphs: List of processed paragraphs

        Returns:
            DocumentMetadata object
        """
        # Calculate word count from all paragraph text
        total_word_count = sum(
            len(para_text.split()) for para_text, _, _, _ in paragraphs
            if para_text.strip()
        )

        return DocumentMetadata(
            filename=file_path.name,
            document_type=DocumentType.WORD,
            file_size=file_path.stat().st_size,
            word_count=total_word_count,
            total_rows=len(paragraphs),  # Use total_rows for paragraph count for consistency
            processing_status=ProcessingStatus.PENDING
        )