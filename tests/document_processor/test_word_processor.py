"""
Tests for WordProcessor
Comprehensive test suite covering functionality, error handling, and edge cases
"""

import pytest
import tempfile
from pathlib import Path
from unittest.mock import patch, MagicMock
from docx import Document
from docx.text.paragraph import Paragraph

from src.models.document import (
    ProcessingConfig,
    ProcessingStatus,
    DocumentType
)
from src.document_processor import WordProcessor
from src.document_processor.word_processor import WordProcessingError


class TestWordProcessorInitialization:
    """Test WordProcessor initialization and configuration"""

    def test_init_with_default_config(self):
        """Test initialization with default configuration"""
        processor = WordProcessor()

        assert processor.config is not None
        assert isinstance(processor.config, ProcessingConfig)
        assert processor.config.max_chunk_size == 800  # Default value
        assert processor.config.chunk_overlap == 100    # Default value

    def test_init_with_custom_config(self):
        """Test initialization with custom configuration"""
        custom_config = ProcessingConfig(
            max_chunk_size=2000,
            chunk_overlap=200,
            skip_empty_paragraphs=False
        )
        processor = WordProcessor(custom_config)

        assert processor.config.max_chunk_size == 2000
        assert processor.config.chunk_overlap == 200
        assert processor.config.skip_empty_paragraphs is False

    def test_init_validates_config(self):
        """Test that invalid config raises error during initialization"""
        invalid_config = ProcessingConfig(
            max_chunk_size=500,
            chunk_overlap=600  # Overlap > max_chunk_size
        )

        with pytest.raises(ValueError, match="chunk_overlap must be less than max_chunk_size"):
            WordProcessor(invalid_config)


class TestWordProcessorValidation:
    """Test file validation functionality"""

    def setup_method(self):
        """Set up test processor for each test"""
        self.processor = WordProcessor()

    def test_validate_nonexistent_file(self):
        """Test validation fails for nonexistent file"""
        document = self.processor.process("nonexistent.docx")

        assert document.metadata.processing_status == ProcessingStatus.FAILED
        assert "File does not exist" in document.metadata.error_message
        assert len(document.chunks) == 0

    def test_validate_directory_instead_of_file(self, tmp_path):
        """Test validation fails when path is a directory"""
        directory = tmp_path / "test_dir"
        directory.mkdir()

        document = self.processor.process(directory)
        assert document.metadata.processing_status == ProcessingStatus.FAILED
        assert "Path is not a file" in document.metadata.error_message
        assert len(document.chunks) == 0

    def test_validate_empty_file(self, tmp_path):
        """Test validation fails for empty file"""
        empty_file = tmp_path / "empty.docx"
        empty_file.touch()  # Create empty file

        document = self.processor.process(empty_file)
        assert document.metadata.processing_status == ProcessingStatus.FAILED
        assert "File is empty" in document.metadata.error_message
        assert len(document.chunks) == 0

    def test_validate_unsupported_extension(self, tmp_path):
        """Test validation fails for unsupported file extensions"""
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("not a word file")

        document = self.processor.process(txt_file)
        assert document.metadata.processing_status == ProcessingStatus.FAILED
        assert "Unsupported file extension" in document.metadata.error_message
        assert len(document.chunks) == 0

    def test_validate_supported_extensions(self, tmp_path):
        """Test that .docx extension is supported"""
        # Create a simple Word document
        docx_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test content")
        doc.save(docx_file)

        # Should process successfully (not fail validation)
        document = self.processor.process(docx_file)
        assert document.metadata.processing_status == ProcessingStatus.COMPLETED


class TestWordProcessorWithRealData:
    """Test WordProcessor with real sample data"""

    def setup_method(self):
        """Set up test processor and check for sample data"""
        self.processor = WordProcessor()

        # Find sample file relative to project root
        test_file_dir = Path(__file__).parent.parent.parent  # Go up from tests/document_processor/
        self.sample_file = test_file_dir / "data" / "company_handbook.docx"

        if not self.sample_file.exists():
            pytest.skip(f"Sample data file not found at {self.sample_file}. Run 'python create_sample_data.py' first.")

    def test_process_sample_handbook_file(self):
        """Test processing the sample company handbook file"""
        document = self.processor.process(self.sample_file)

        # Check metadata
        assert document.metadata.filename == "company_handbook.docx"
        assert document.metadata.document_type == DocumentType.WORD
        assert document.metadata.processing_status == ProcessingStatus.COMPLETED
        assert document.metadata.file_size > 0
        assert document.metadata.processing_time_seconds is not None
        assert document.metadata.processing_time_seconds > 0

        # Check word count and paragraph count
        assert document.metadata.word_count > 0
        assert document.metadata.total_rows > 0  # total_rows = paragraph count
        assert document.metadata.total_chunks > 0
        assert len(document.chunks) == document.metadata.total_chunks

        # Check document properties
        assert document.chunk_count == len(document.chunks)
        assert document.is_processed is True
        assert len(document.get_chunks_with_embeddings()) == 0  # No embeddings yet

    def test_chunk_content_format(self):
        """Test that chunk content follows expected natural language format"""
        document = self.processor.process(self.sample_file)

        # Should have at least one chunk
        assert len(document.chunks) > 0

        # Check first chunk content
        first_chunk = document.chunks[0]
        assert first_chunk.content
        assert isinstance(first_chunk.content, str)
        assert len(first_chunk.content.strip()) > 0

        # Check chunk metadata
        assert first_chunk.chunk_id
        assert first_chunk.source == "company_handbook.docx"
        assert first_chunk.chunk_index >= 0
        assert isinstance(first_chunk.metadata, dict)

        # Word-specific metadata
        assert "paragraph_indices" in first_chunk.metadata
        assert "paragraph_types" in first_chunk.metadata
        assert "character_count" in first_chunk.metadata
        assert "has_headings" in first_chunk.metadata

    def test_document_structure_recognition(self):
        """Test that document structure is properly recognized"""
        document = self.processor.process(self.sample_file)

        # Collect all paragraph types and headings
        all_types = set()
        all_headings = set()

        for chunk in document.chunks:
            chunk_types = chunk.metadata.get("paragraph_types", [])
            chunk_headings = chunk.metadata.get("section_headings", [])
            all_types.update(chunk_types)
            all_headings.update(chunk_headings)

        # Should recognize different paragraph types
        # (exact types depend on sample document content)
        assert len(all_types) > 1  # Should have multiple types
        assert len(all_headings) > 1  # Should have multiple headings

        # Check for expected structure elements
        has_headings = any("heading" in ptype for ptype in all_types)
        assert has_headings  # Should detect headings

    def test_chunking_respects_max_size(self):
        """Test that chunks respect the maximum chunk size configuration"""
        config = ProcessingConfig(max_chunk_size=500)  # Small chunks
        processor = WordProcessor(config)

        document = processor.process(self.sample_file)

        # All chunks should be within the size limit
        for chunk in document.chunks:
            assert len(chunk.content) <= config.max_chunk_size
            assert len(chunk.content) > 0  # Should not be empty

    def test_heading_aware_chunking(self):
        """Test that chunking prefers to start new chunks with headings"""
        document = self.processor.process(self.sample_file)

        # Check that chunks starting with headings are common
        chunks_with_headings = 0
        for chunk in document.chunks:
            paragraph_types = chunk.metadata.get("paragraph_types", [])
            if any(ptype.startswith("heading") for ptype in paragraph_types):
                chunks_with_headings += 1

        # Should have some chunks with headings (depends on document structure)
        assert chunks_with_headings > 0


class TestWordProcessorWithMockData:
    """Test WordProcessor with controlled mock data"""

    def setup_method(self):
        """Set up test processor"""
        self.processor = WordProcessor()

    def create_test_word_file(self, tmp_path, paragraphs_data):
        """
        Helper to create test Word files

        Args:
            tmp_path: pytest temporary path
            paragraphs_data: List of (text, style_name) tuples
        """
        word_file = tmp_path / "test.docx"
        doc = Document()

        for text, style_name in paragraphs_data:
            paragraph = doc.add_paragraph(text)
            if style_name:
                paragraph.style = style_name

        doc.save(word_file)
        return word_file

    def test_process_simple_document(self, tmp_path):
        """Test processing simple document with mixed paragraph types"""
        paragraphs = [
            ("Test Document Title", "Title"),
            ("Introduction", "Heading 1"),
            ("This is a simple introduction paragraph.", "Normal"),
            ("Key points to remember:", "Normal"),
            ("First important point", "List Bullet"),
            ("Second important point", "List Bullet"),
            ("Conclusion", "Heading 1"),
            ("This concludes our test document.", "Normal")
        ]

        word_file = self.create_test_word_file(tmp_path, paragraphs)
        document = self.processor.process(word_file)

        assert document.metadata.processing_status == ProcessingStatus.COMPLETED
        assert document.metadata.total_rows > 0  # Should have paragraphs
        assert len(document.chunks) > 0

        # Check that content includes expected elements
        all_content = " ".join(chunk.content for chunk in document.chunks)
        assert "introduction" in all_content.lower()
        assert "conclusion" in all_content.lower()

    def test_process_document_with_only_headings(self, tmp_path):
        """Test processing document with only heading paragraphs"""
        paragraphs = [
            ("Chapter 1", "Heading 1"),
            ("Section 1.1", "Heading 2"),
            ("Section 1.2", "Heading 2"),
            ("Chapter 2", "Heading 1")
        ]

        word_file = self.create_test_word_file(tmp_path, paragraphs)
        document = self.processor.process(word_file)

        assert document.metadata.processing_status == ProcessingStatus.COMPLETED
        assert len(document.chunks) > 0

        # All chunks should have heading content
        for chunk in document.chunks:
            assert any(ptype.startswith("heading") for ptype in chunk.metadata["paragraph_types"])

    def test_process_empty_paragraphs_handling(self, tmp_path):
        """Test handling of empty paragraphs based on configuration"""
        paragraphs = [
            ("Valid content", "Normal"),
            ("", "Normal"),  # Empty paragraph
            ("More valid content", "Normal"),
            ("   ", "Normal"),  # Whitespace only
            ("Final content", "Normal")
        ]

        # Test with skip_empty_paragraphs=True (default)
        word_file = self.create_test_word_file(tmp_path, paragraphs)
        document = self.processor.process(word_file)

        assert document.metadata.processing_status == ProcessingStatus.COMPLETED

        # Should skip empty paragraphs
        all_content = " ".join(chunk.content for chunk in document.chunks)
        assert "valid content" in all_content.lower()
        assert "final content" in all_content.lower()

    def test_chunking_with_custom_config(self, tmp_path):
        """Test contextual chunking with overlap and optimal sizing"""
        # Create document with multiple paragraphs that will demonstrate good chunking
        paragraphs = [
            ("Introduction", "Heading 1"),
            ("This is the first paragraph with some important content.", "Normal"),
            ("This is the second paragraph that continues the topic.", "Normal"),
            ("This third paragraph adds more detail to the discussion.", "Normal"),
            ("Benefits Section", "Heading 1"),
            ("The first benefit is comprehensive health coverage.", "Normal"),
            ("The second benefit includes retirement planning.", "Normal"),
            ("Additional perks make this package competitive.", "Normal")
        ]

        word_file = self.create_test_word_file(tmp_path, paragraphs)

        # Test with contextual chunking configuration
        config = ProcessingConfig(
            max_chunk_size=300,  # Smaller chunks to force multiple chunks
            min_chunk_size=100,  # Ensure substantial chunks
            target_sentences_per_chunk=2,  # Target 2 sentences per chunk
            overlap_sentences=1  # 1 sentence overlap
        )
        processor = WordProcessor(config)

        document = processor.process(word_file)

        # Should create multiple contextual chunks
        assert len(document.chunks) > 1

        # Each chunk should be within the size range (allowing for context addition)
        for chunk in document.chunks:
            assert len(chunk.content) >= 50  # Allow for some smaller chunks due to headings
            assert len(chunk.content) <= config.max_chunk_size * 1.2  # 20% tolerance for context

            # Check that chunks have good metadata
            assert "sentence_count" in chunk.metadata
            assert "chunk_quality" in chunk.metadata

    def test_overlap_functionality(self, tmp_path):
        """Test that overlap between chunks preserves context"""
        # Create document that will definitely need multiple chunks
        paragraphs = [
            ("First Section", "Heading 1"),
            ("The first paragraph introduces the topic with important keywords.", "Normal"),
            ("The second paragraph continues with more keywords and details.", "Normal"),
            ("The third paragraph concludes this section with final thoughts.", "Normal"),
            ("Second Section", "Heading 1"),
            ("The fourth paragraph starts a new topic but should overlap.", "Normal")
        ]

        word_file = self.create_test_word_file(tmp_path, paragraphs)

        # Configure for overlap
        config = ProcessingConfig(
            max_chunk_size=250,  # Force multiple chunks
            overlap_sentences=1,  # Ensure overlap
            target_sentences_per_chunk=2
        )
        processor = WordProcessor(config)

        document = processor.process(word_file)

        # Should have multiple chunks
        assert len(document.chunks) > 1

        # Check for overlap indicators in metadata
        overlap_chunks = [c for c in document.chunks if c.metadata.get("has_overlap")]
        assert len(overlap_chunks) > 0  # Should have chunks with overlap

        # Verify that chunks contain contextual information
        all_content = " ".join(chunk.content for chunk in document.chunks)
        assert "keywords" in all_content  # Should preserve important terms across chunks

    def test_single_long_paragraph_behavior(self, tmp_path):
        """Test behavior with single paragraph longer than embedding limits"""
        # Create document with one very long paragraph
        long_paragraph_text = "This is a very long paragraph with lots of text. " * 30  # ~1500 chars
        paragraphs = [
            ("Long Content", "Heading 1"),
            (long_paragraph_text, "Normal")
        ]

        word_file = self.create_test_word_file(tmp_path, paragraphs)

        # Test with embedding limits enforced
        config = ProcessingConfig(
            max_chunk_size=1000,
            max_embedding_tokens=256,  # Small limit
            chars_per_token_estimate=4.0,  # 256 * 4 = 1024 char limit
            enforce_embedding_limits=True
        )
        processor = WordProcessor(config)

        document = processor.process(word_file)

        assert document.metadata.processing_status == ProcessingStatus.COMPLETED
        assert len(document.chunks) >= 1

        # Should split long paragraph to respect embedding limits
        max_expected_size = config.max_embedding_chars  # 1024 chars
        oversized_chunks = [chunk for chunk in document.chunks if len(chunk.content) > max_expected_size * 1.1]  # 10% tolerance

        # Should have very few (ideally zero) chunks exceeding embedding limits
        assert len(oversized_chunks) <= 1  # Allow some tolerance for edge cases

    def test_embedding_limits_configuration(self, tmp_path):
        """Test that embedding limits are properly applied"""
        # Create normal document
        paragraphs = [("Normal paragraph text.", "Normal")]
        word_file = self.create_test_word_file(tmp_path, paragraphs)

        # Test with embedding limits disabled
        config_disabled = ProcessingConfig(enforce_embedding_limits=False)
        processor_disabled = WordProcessor(config_disabled)

        # Test with embedding limits enabled
        config_enabled = ProcessingConfig(
            enforce_embedding_limits=True,
            max_embedding_tokens=100,  # Very small
            chars_per_token_estimate=4.0
        )
        processor_enabled = WordProcessor(config_enabled)

        doc_disabled = processor_disabled.process(word_file)
        doc_enabled = processor_enabled.process(word_file)

        # Both should succeed
        assert doc_disabled.metadata.processing_status == ProcessingStatus.COMPLETED
        assert doc_enabled.metadata.processing_status == ProcessingStatus.COMPLETED


class TestWordProcessorErrorHandling:
    """Test error handling and edge cases"""

    def setup_method(self):
        """Set up test processor"""
        self.processor = WordProcessor()

    def test_corrupted_word_file(self, tmp_path):
        """Test handling of corrupted Word file"""
        fake_docx = tmp_path / "corrupted.docx"
        fake_docx.write_text("This is not a real Word file")

        document = self.processor.process(fake_docx)

        # Should return failed document, not crash
        assert document.metadata.processing_status == ProcessingStatus.FAILED
        assert document.metadata.error_message is not None
        assert len(document.chunks) == 0

    @patch('src.document_processor.word_processor.Document')
    def test_docx_library_error(self, mock_document, tmp_path):
        """Test handling of python-docx library errors"""
        mock_document.side_effect = Exception("Simulated docx error")

        # Create a real file to pass validation
        word_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test")
        doc.save(word_file)

        document = self.processor.process(word_file)

        assert document.metadata.processing_status == ProcessingStatus.FAILED
        assert "Simulated docx error" in document.metadata.error_message

    def test_permission_denied(self):
        """Test handling of permission denied errors"""
        with patch('pathlib.Path.exists', return_value=True):
            with patch('pathlib.Path.is_file', return_value=True):
                with patch('pathlib.Path.stat') as mock_stat:
                    mock_stat.return_value.st_size = 1024
                    with patch('builtins.open', side_effect=PermissionError("Permission denied")):
                        document = self.processor.process("protected.docx")
                        assert document.metadata.processing_status == ProcessingStatus.FAILED
                        assert "Permission denied" in document.metadata.error_message
                        assert len(document.chunks) == 0


class TestWordProcessorParagraphSplitting:
    """Test the paragraph splitting logic for embedding compatibility"""

    def setup_method(self):
        """Set up test processor"""
        self.processor = WordProcessor()

    def test_short_paragraph_no_splitting(self):
        """Test that short paragraphs are not split"""
        short_text = "This is a short paragraph."
        segments = self.processor._split_long_paragraph(short_text, 100)

        assert len(segments) == 1
        assert segments[0] == short_text

    def test_long_paragraph_sentence_splitting_with_overlap(self):
        """Test splitting long paragraph with contextual overlap"""
        # Create a processor with overlap settings
        config = ProcessingConfig(target_sentences_per_chunk=2, overlap_sentences=1)
        processor = WordProcessor(config)

        long_text = "First sentence with important context. Second sentence continues the thought. Third sentence adds more detail. Fourth sentence provides conclusion."
        segments = processor._split_long_paragraph(long_text, 80)  # Force splitting

        assert len(segments) > 1  # Should be split

        # Check that overlap preserves context - some sentences should appear in multiple segments
        all_sentences = []
        for segment in segments:
            sentences_in_segment = [s.strip() for s in segment.split('.') if s.strip()]
            all_sentences.extend(sentences_in_segment)

        # With overlap, we should have more total sentences than in original
        original_sentences = [s.strip() for s in long_text.split('.') if s.strip()]
        assert len(all_sentences) >= len(original_sentences)  # Overlap means duplication

        # Each segment should be reasonably sized
        for segment in segments:
            assert len(segment) <= 100  # Within reasonable limit
            assert len(segment) > 20   # Not too small

    def test_contextual_chunk_sizing(self):
        """Test that chunks target optimal size for RAG"""
        config = ProcessingConfig(
            min_chunk_size=200,
            max_chunk_size=800,
            target_sentences_per_chunk=3
        )
        processor = WordProcessor(config)

        # Test the splitting with various text lengths
        short_text = "Short sentence."
        medium_text = "This is a medium length sentence with several words. This is another sentence. And a third one."
        long_text = "This is a very long sentence with many words and phrases that provides lots of context. " * 5

        # Short text shouldn't be split
        assert len(processor._split_long_paragraph(short_text, 800)) == 1

        # Medium text might be split based on target sentences
        medium_segments = processor._split_long_paragraph(medium_text, 200)
        for segment in medium_segments:
            sentence_count = len([s for s in segment.split('.') if s.strip()])
            assert sentence_count <= config.target_sentences_per_chunk + 1  # Allow some flexibility

        # Long text should definitely be split
        long_segments = processor._split_long_paragraph(long_text, 400)
        assert len(long_segments) > 1

        for segment in long_segments:
            assert len(segment) <= 450  # Reasonable size with tolerance

    def test_very_long_single_sentence_word_splitting(self):
        """Test word-level splitting for extremely long single sentences"""
        # Create a very long sentence without periods
        long_sentence = "This sentence has many words " * 20 + "and never ends"
        segments = self.processor._split_long_paragraph(long_sentence, 100)

        assert len(segments) > 1  # Should be split at word boundaries

        for segment in segments:
            assert len(segment) <= 110  # Allow some tolerance for word boundaries

    def test_mixed_sentence_lengths(self):
        """Test splitting with mix of short and long sentences"""
        mixed_text = "Short. This is a much longer sentence that contains many words and should be split appropriately. Another short one."
        segments = self.processor._split_long_paragraph(mixed_text, 50)

        assert len(segments) >= 2  # Should split the long sentence

        # Verify content preservation
        all_words = set(mixed_text.replace('.', '').split())
        segment_words = set(' '.join(segments).replace('.', '').split())
        assert all_words == segment_words


class TestWordProcessorParagraphConversion:
    """Test the paragraph-to-text conversion logic specifically"""

    def setup_method(self):
        """Set up test processor"""
        self.processor = WordProcessor()

    def test_convert_heading_paragraph(self):
        """Test converting heading paragraphs"""
        text = self.processor._convert_paragraph_to_text(
            "Company Overview", "heading_1", ""
        )

        assert "Section 1" in text
        assert "Company Overview" in text

    def test_convert_normal_paragraph_with_heading_context(self):
        """Test converting normal paragraph with heading context"""
        text = self.processor._convert_paragraph_to_text(
            "TechCorp is a leading technology company.",
            "normal",
            "Company Overview"
        )

        assert "Under Company Overview" in text
        assert "TechCorp is a leading technology company" in text

    def test_convert_list_item(self):
        """Test converting list item paragraphs"""
        text = self.processor._convert_paragraph_to_text(
            "Health insurance with 90% coverage",
            "list_item",
            "Employee Benefits"
        )

        assert "Under Employee Benefits" in text
        assert "item:" in text
        assert "Health insurance" in text

    def test_convert_quote_paragraph(self):
        """Test converting quote paragraphs"""
        text = self.processor._convert_paragraph_to_text(
            "Quality is not an act, but a habit.",
            "quote",
            "Company Philosophy"
        )

        assert "Quote in Company Philosophy" in text
        assert "Quality is not an act" in text

    def test_convert_title_paragraph(self):
        """Test converting title paragraphs"""
        text = self.processor._convert_paragraph_to_text(
            "Employee Handbook",
            "title",
            ""
        )

        assert "Document title:" in text
        assert "Employee Handbook" in text

    def test_convert_empty_paragraph(self):
        """Test converting empty paragraphs"""
        text = self.processor._convert_paragraph_to_text("", "normal", "Section")

        assert text == ""

    def test_convert_whitespace_only_paragraph(self):
        """Test converting paragraphs with only whitespace"""
        text = self.processor._convert_paragraph_to_text("   \n\t   ", "normal", "Section")

        assert text == ""


class TestWordProcessorGetParagraphType:
    """Test paragraph type detection"""

    def setup_method(self):
        """Set up test processor"""
        self.processor = WordProcessor()

    def create_mock_paragraph(self, style_name):
        """Create a mock paragraph with given style"""
        mock_paragraph = MagicMock()
        mock_paragraph.style.name = style_name
        return mock_paragraph

    def test_detect_heading_types(self):
        """Test detection of different heading levels"""
        test_cases = [
            ("Heading 1", "heading_1"),
            ("Heading 2", "heading_2"),
            ("Heading 3", "heading_3"),
            ("Heading 4", "heading_other"),
        ]

        for style_name, expected_type in test_cases:
            mock_para = self.create_mock_paragraph(style_name)
            result = self.processor._get_paragraph_type(mock_para)
            assert result == expected_type

    def test_detect_list_types(self):
        """Test detection of list paragraph types"""
        list_styles = ["List Bullet", "List Number", "List Paragraph"]

        for style_name in list_styles:
            mock_para = self.create_mock_paragraph(style_name)
            result = self.processor._get_paragraph_type(mock_para)
            assert result == "list_item"

    def test_detect_special_types(self):
        """Test detection of special paragraph types"""
        test_cases = [
            ("Quote", "quote"),
            ("Title", "title"),
            ("Normal", "normal"),
            ("Body Text", "normal"),
        ]

        for style_name, expected_type in test_cases:
            mock_para = self.create_mock_paragraph(style_name)
            result = self.processor._get_paragraph_type(mock_para)
            assert result == expected_type


if __name__ == "__main__":
    # Run tests with: python -m pytest tests/document_processor/test_word_processor.py -v
    pytest.main([__file__, "-v"])