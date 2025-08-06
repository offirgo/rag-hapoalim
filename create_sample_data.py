"""
Script to create sample data files for RAG system testing
Run this once to generate sample Excel and Word files
"""

import pandas as pd
from docx import Document
import os

def create_sample_excel():
    """Create sample employee data Excel file"""

    # Sample employee data
    employees = [
        {"Employee_ID": "EMP001", "Name": "John Smith", "Department": "Engineering",
         "Role": "Senior Software Developer", "Salary": 95000, "Location": "New York",
         "Skills": "Python, React, AWS", "Years_Experience": 8, "Manager": "Sarah Johnson"},

        {"Employee_ID": "EMP002", "Name": "Emily Chen", "Department": "Marketing",
         "Role": "Marketing Manager", "Salary": 75000, "Location": "San Francisco",
         "Skills": "Digital Marketing, Analytics, SEO", "Years_Experience": 5, "Manager": "Mike Wilson"},

        {"Employee_ID": "EMP003", "Name": "Michael Rodriguez", "Department": "Engineering",
         "Role": "DevOps Engineer", "Salary": 88000, "Location": "Austin",
         "Skills": "Docker, Kubernetes, Jenkins", "Years_Experience": 6, "Manager": "Sarah Johnson"},

        {"Employee_ID": "EMP004", "Name": "Sarah Johnson", "Department": "Engineering",
         "Role": "Engineering Manager", "Salary": 120000, "Location": "New York",
         "Skills": "Team Leadership, Architecture, Python", "Years_Experience": 12, "Manager": "CEO"},

        {"Employee_ID": "EMP005", "Name": "David Kim", "Department": "Sales",
         "Role": "Sales Representative", "Salary": 65000, "Location": "Chicago",
         "Skills": "CRM, Negotiation, Client Relations", "Years_Experience": 3, "Manager": "Lisa Brown"},

        {"Employee_ID": "EMP006", "Name": "Lisa Brown", "Department": "Sales",
         "Role": "Sales Director", "Salary": 105000, "Location": "Chicago",
         "Skills": "Sales Strategy, Team Management", "Years_Experience": 10, "Manager": "CEO"},

        {"Employee_ID": "EMP007", "Name": "Amanda Taylor", "Department": "HR",
         "Role": "HR Specialist", "Salary": 58000, "Location": "New York",
         "Skills": "Recruitment, Employee Relations", "Years_Experience": 4, "Manager": "Robert Davis"},

        {"Employee_ID": "EMP008", "Name": "Robert Davis", "Department": "HR",
         "Role": "HR Director", "Salary": 95000, "Location": "New York",
         "Skills": "HR Strategy, Compliance, Leadership", "Years_Experience": 15, "Manager": "CEO"}
    ]

    df = pd.DataFrame(employees)

    # Create Excel file with formatting
    with pd.ExcelWriter('data/sample_employees.xlsx', engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Employees', index=False)

        # Add a second sheet with department summary
        dept_summary = df.groupby('Department').agg({
            'Name': 'count',
            'Salary': 'mean',
            'Years_Experience': 'mean'
        }).round(2)
        dept_summary.columns = ['Employee_Count', 'Average_Salary', 'Average_Experience']
        dept_summary.to_excel(writer, sheet_name='Department_Summary')

    print("✅ Created data/sample_employees.xlsx")

def create_sample_word_doc():
    """Create sample company handbook Word document"""

    doc = Document()

    # Title
    title = doc.add_heading('TechCorp Employee Handbook', 0)

    # Section 1: Company Overview
    doc.add_heading('1. Company Overview', level=1)
    doc.add_paragraph(
        'TechCorp is a leading technology company founded in 2015, specializing in '
        'cloud-based software solutions. Our mission is to empower businesses through '
        'innovative technology and exceptional customer service. We are headquartered '
        'in New York with offices in San Francisco, Austin, and Chicago.'
    )

    # Section 2: Work Schedule
    doc.add_heading('2. Work Schedule and Remote Work Policy', level=1)
    doc.add_paragraph(
        'Standard work hours are Monday through Friday, 9:00 AM to 5:00 PM local time. '
        'TechCorp supports flexible working arrangements and remote work options. '
        'Employees may work remotely up to 3 days per week with manager approval. '
        'Core collaboration hours are 10:00 AM to 3:00 PM when all team members '
        'should be available for meetings and collaboration.'
    )

    # Section 3: Benefits
    doc.add_heading('3. Employee Benefits', level=1)
    doc.add_paragraph(
        'TechCorp provides comprehensive benefits including:'
    )
    benefits_list = [
        'Health insurance with 90% company coverage',
        'Dental and vision insurance',
        '401(k) retirement plan with 4% company match',
        '20 days paid time off annually',
        '10 paid holidays per year',
        '$2,000 annual professional development budget',
        'Flexible spending accounts for health and dependent care'
    ]
    for benefit in benefits_list:
        doc.add_paragraph(benefit, style='List Bullet')

    # Section 4: Code of Conduct
    doc.add_heading('4. Code of Conduct', level=1)
    doc.add_paragraph(
        'All employees are expected to maintain high standards of professional conduct. '
        'This includes treating colleagues with respect, maintaining confidentiality '
        'of sensitive information, and adhering to company policies. Harassment, '
        'discrimination, or unethical behavior will not be tolerated.'
    )

    # Section 5: IT and Security
    doc.add_heading('5. IT and Security Policies', level=1)
    doc.add_paragraph(
        'Employees must use company-approved software and follow security protocols. '
        'All devices must have up-to-date antivirus software and operating system patches. '
        'Use of personal devices for work requires IT department approval and security '
        'configuration. Passwords must be complex and changed every 90 days. '
        'Two-factor authentication is required for all company systems.'
    )

    # Section 6: Performance Reviews
    doc.add_heading('6. Performance Management', level=1)
    doc.add_paragraph(
        'Performance reviews are conducted annually in January. Employees set goals '
        'with their managers and receive feedback quarterly. Career development '
        'discussions are encouraged, and the company supports internal promotions. '
        'Performance improvement plans may be implemented when necessary to help '
        'employees meet expectations.'
    )

    # Section 7: Emergency Procedures
    doc.add_heading('7. Emergency Procedures', level=1)
    doc.add_paragraph(
        'In case of emergency, employees should follow building evacuation procedures '
        'and gather at designated meeting points. For workplace injuries, contact '
        'security immediately at extension 911. For IT emergencies or security '
        'incidents, contact the IT helpdesk at ext. 4357 or it-emergency@techcorp.com.'
    )

    # Save the document
    doc.save('data/company_handbook.docx')
    print("✅ Created data/company_handbook.docx")

def main():
    # Create data directory if it doesn't exist
    os.makedirs('data', exist_ok=True)

    # Create sample files
    create_sample_excel()
    create_sample_word_doc()

    print("\n🎉 Sample data files created successfully!")
    print("📁 Files created:")
    print("   - data/sample_employees.xlsx (8 employees across 4 departments)")
    print("   - data/company_handbook.docx (7 sections of company policies)")
    print("\n💡 You can now use these files to test your RAG system!")

if __name__ == "__main__":
    main()